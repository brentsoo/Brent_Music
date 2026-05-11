import streamlit as st
from streamlit_gsheets import GSheetsConnection
import pandas as pd
import json
from datetime import datetime, date, timedelta, time as dtime
import time
import requests

# ─────────────────────────────────────────────
# 1. PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(page_title="Brent Music", layout="wide", page_icon="🎹")

# ─────────────────────────────────────────────
# 2. LOGIN
# ─────────────────────────────────────────────
def check_login():
    if "logged_in" not in st.session_state:
        st.session_state.logged_in = False
    if not st.session_state.logged_in:
        st.title("🎹 Brent Music 登录")
        st.write("")
        username = st.text_input("用户名")
        password = st.text_input("密码", type="password")
        if st.button("登录"):
            users = st.secrets["passwords"]
            if username in users and users[username] == password:
                st.session_state.logged_in = True
                st.session_state.username = username
                st.session_state.role = st.secrets["roles"][username]
                st.rerun()
            else:
                st.error("❌ 用户名或密码错误")
        st.stop()

check_login()
role     = st.session_state.role
username = st.session_state.username

# ─────────────────────────────────────────────
# 3. CONSTANTS
# ─────────────────────────────────────────────
conn = st.connection("gsheets", type=GSheetsConnection)
DAYS_OPTIONS = ["周一","周二","周三","周四","周五","周六","周日"]
DAY_MAP      = {"周一":0,"周二":1,"周三":2,"周四":3,"周五":4,"周六":5,"周日":6}
INSTRUMENTS  = ["钢琴","小提琴","吉他","大提琴","其他"]

# ─────────────────────────────────────────────
# 4. DATA HELPERS
# ─────────────────────────────────────────────
def safe_json(val):
    if pd.isna(val) or str(val).strip() == "":
        return []
    try:
        return json.loads(val)
    except:
        return []

def load_data():
    try:
        df = conn.read(worksheet="工作表1", ttl=0)
        if df.empty:
            return {}
        db = {}
        for _, row in df.iterrows():
            sid = str(row['student_id'])
            db[sid] = {
                "name":                str(row['name']),
                "instrument":          str(row['instrument']),
                "grade":               str(row['grade']),
                "phone":               str(row.get('phone', '')),
                "email":               str(row.get('email', '')),
                "replacement_credits": int(row['replacement_credits']) if pd.notna(row['replacement_credits']) else 0,
                "replacement_minutes": int(row['replacement_minutes']) if pd.notna(row.get('replacement_minutes')) else 0,
                "history":             safe_json(row['history']),
                "schedule":            safe_json(row.get('schedule', '')),
            }
        return db
    except:
        return {}

def save_data(db):
    rows = []
    for sid, info in db.items():
        rows.append({
            "student_id":          sid,
            "name":                info['name'],
            "instrument":          info['instrument'],
            "grade":               info['grade'],
            "phone":               info.get('phone', ''),
            "email":               info.get('email', ''),
            "replacement_credits": info['replacement_credits'],
            "replacement_minutes": info.get('replacement_minutes', 0),
            "history":             json.dumps(info['history'],  ensure_ascii=False),
            "schedule":            json.dumps(info.get('schedule',[]), ensure_ascii=False),
        })
    conn.update(worksheet="工作表1", data=pd.DataFrame(rows))
    st.cache_data.clear()

def load_pending():
    try:
        df = conn.read(worksheet="待审批", ttl=0)
        if df.empty:
            return []
        return [{
            "id":           str(r['id']),
            "requested_by": str(r['requested_by']),
            "action":       str(r['action']),
            "data":         json.loads(str(r['data'])),
            "timestamp":    str(r['timestamp']),
        } for _, r in df.iterrows()]
    except:
        return []

def save_pending(lst):
    if not lst:
        df = pd.DataFrame(columns=["id","requested_by","action","data","timestamp"])
    else:
        df = pd.DataFrame([{
            "id": p['id'], "requested_by": p['requested_by'],
            "action": p['action'],
            "data": json.dumps(p['data'], ensure_ascii=False),
            "timestamp": p['timestamp'],
        } for p in lst])
    conn.update(worksheet="待审批", data=df)
    st.cache_data.clear()

def submit_pending(action, data):
    lst = load_pending()
    lst.append({
        "id":           str(int(time.time())),
        "requested_by": username,
        "action":       action,
        "data":         data,
        "timestamp":    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
    })
    save_pending(lst)

def load_todos():
    try:
        df = conn.read(worksheet="ToDoList", ttl=0)
        if df.empty:
            return []
        return [{
            "id":        str(r['id']),
            "task":      str(r['task']),
            "done":      str(r['done']) == "True",
            "created_by":str(r['created_by']),
            "timestamp": str(r['timestamp']),
        } for _, r in df.iterrows()]
    except:
        return []

def save_todos(lst):
    if not lst:
        df = pd.DataFrame(columns=["id","task","done","created_by","timestamp"])
    else:
        df = pd.DataFrame([{
            "id": t['id'], "task": t['task'],
            "done": str(t['done']), "created_by": t['created_by'],
            "timestamp": t['timestamp'],
        } for t in lst])
    conn.update(worksheet="ToDoList", data=df)
    st.cache_data.clear()

def load_events():
    try:
        df = conn.read(worksheet="Events", ttl=0)
        if df.empty:
            return []
        return [{
            "id":          str(r['id']),
            "student_id":  str(r['student_id']),
            "student_name":str(r['student_name']),
            "event_type":  str(r['event_type']),
            "event_name":  str(r['event_name']),
            "event_date":  str(r['event_date']),
            "notes":       str(r['notes']),
            "created_by":  str(r['created_by']),
        } for _, r in df.iterrows()]
    except:
        return []

def save_events(lst):
    if not lst:
        df = pd.DataFrame(columns=["id","student_id","student_name","event_type","event_name","event_date","notes","created_by"])
    else:
        df = pd.DataFrame(lst)
    conn.update(worksheet="Events", data=df)
    st.cache_data.clear()

# ─────────────────────────────────────────────
# 5. GEMINI AI HELPER
# ─────────────────────────────────────────────
def call_gemini(prompt: str) -> str:
    try:
        api_key = st.secrets["gemini"]["api_key"]
        model_name = "gemini-2.5-flash"
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
        headers = {"Content-Type": "application/json"}
        payload = {"contents": [{"parts": [{"text": prompt}]}]}
        resp = requests.post(url, headers=headers, json=payload, timeout=30)
        data = resp.json()
        if "error" in data:
            return f"Gemini 错误：{data['error'].get('message', str(data['error']))}"
        if "candidates" not in data:
            return f"Gemini 返回异常：{json.dumps(data)}"
        candidate = data["candidates"][0]
        if candidate.get("finishReason") == "SAFETY":
            return "AI 内容被安全过滤。"
        return candidate["content"]["parts"][0]["text"]
    except Exception as e:
        return f"连接失败: {e}"

def build_prompt(sid, info, lang):
    h         = info["history"]
    total     = len(h)
    attended  = sum(1 for x in h if x.get("status") == "Attended")
    cancelled = sum(1 for x in h if "Cancelled" in x.get("status",""))
    temp      = sum(1 for x in h if x.get("status") == "Temporary Stop")
    replaced  = sum(1 for x in h if "Replace" in x.get("status",""))
    rate      = f"{attended/total*100:.1f}%" if total else "N/A"
    recent    = "\n".join(
        f"- {x.get('date','')}: {x.get('remarks','')}"
        for x in h[-5:]
        if x.get('remarks') and 'N/A' not in str(x.get('remarks',''))
    ) or ("No remarks." if lang == "English" else "暂无备注。")
    pending_min = info.get('replacement_minutes', 0)

    if lang == "English":
        return f"""You are a professional music teacher assistant. Write a concise student progress report in English.
Student: {info['name']} | Instrument: {info['instrument']} | Grade: {info['grade']}
Lessons: {total} total | Attended: {attended} | Cancelled: {cancelled} | Temp Stop: {temp} | Replacements done: {replaced}
Attendance rate: {rate} | Pending replacement lessons: {info['replacement_credits']} | Pending replacement minutes: {pending_min} min
Recent remarks:\n{recent}
Write 3-5 sentences covering attendance, learning progress, and a brief recommendation."""
    else:
        return f"""你是一位专业音乐教师助理。请用中文写一份简洁的学生进度报告。
姓名：{info['name']} | 乐器：{info['instrument']} | 等级：{info['grade']}
总课次：{total} | 出席：{attended} | 取消：{cancelled} | 暂停：{temp} | 已补：{replaced}
出席率：{rate} | 待补课次数：{info['replacement_credits']} | 待补课时长：{pending_min} 分钟
最近备注：\n{recent}
请写3-5句话，涵盖出席情况、学习进度及简短建议。"""

# ─────────────────────────────────────────────
# 6. LOAD DATA & SIDEBAR
# ─────────────────────────────────────────────
students_db = load_data()

st.sidebar.title("🎹 Brent Music")
st.sidebar.write(f"👤 **{username}**")
st.sidebar.caption({"owner":"🔑 Owner","editor":"✏️ Editor","viewer":"👁️ Viewer"}[role])
if st.sidebar.button("退出登录"):
    st.session_state.logged_in = False
    st.rerun()
st.sidebar.divider()

if role == "owner":
    menu_options = ["添加学生","考勤登记","补课安排","报告查询","管理学生信息","报表中心","课程日历","📝 To Do List","🎪 Events","⚠️ 审批请求"]
elif role == "editor":
    menu_options = ["考勤登记","补课安排","报告查询","报表中心","课程日历","📝 To Do List","🎪 Events","申请修改学生信息"]
else:
    menu_options = ["报告查询","报表中心","课程日历"]

menu = st.sidebar.selectbox("📅 功能菜单", menu_options)

# ─────────────────────────────────────────────
# A. 添加学生
# ─────────────────────────────────────────────
if menu == "添加学生":
    st.title("➕ 新生入库")
    with st.form("add_student"):
        col1, col2 = st.columns(2)
        with col1:
            sid   = st.text_input("学生 ID（唯一）")
            name  = st.text_input("学生姓名")
            inst  = st.selectbox("乐器", INSTRUMENTS)
            grade = st.text_input("当前等级")
            phone = st.text_input("联络号码")
            email = st.text_input("电子邮箱")
        with col2:
            st.write("**📅 固定上课时间**")
            duration = st.number_input("每节课时长（分钟）", min_value=15, max_value=180, value=60, step=15)
            num_slots = st.number_input("每周上课次数", min_value=1, max_value=7, value=1, step=1)
            schedule_slots = []
            for i in range(int(num_slots)):
                c1, c2 = st.columns(2)
                with c1:
                    day = st.selectbox(f"第{i+1}节 – 周几", DAYS_OPTIONS, key=f"day_{i}")
                with c2:
                    start_t = st.time_input(f"第{i+1}节 – 开始时间", key=f"time_{i}")
                # Auto-calculate end time
                start_dt = datetime.combine(date.today(), start_t)
                end_dt   = start_dt + timedelta(minutes=int(duration))
                end_t    = end_dt.time()
                st.caption(f"↳ 结束时间：{end_t.strftime('%H:%M')}")
                schedule_slots.append({
                    "day": day,
                    "start": str(start_t),
                    "end":   end_t.strftime("%H:%M:%S"),
                    "duration_min": int(duration),
                })
        if st.form_submit_button("确认添加"):
            if sid and name:
                if sid not in students_db:
                    students_db[sid] = {
                        "name": name, "instrument": inst, "grade": grade,
                        "phone": phone, "email": email,
                        "history": [], "replacement_credits": 0,
                        "replacement_minutes": 0,
                        "schedule": schedule_slots,
                    }
                    save_data(students_db)
                    st.success(f"✅ {name} 已成功加入！")
                    st.rerun()
                else:
                    st.error("❌ 该 ID 已存在。")
            else:
                st.error("❌ 请填入 ID 和姓名")

# ─────────────────────────────────────────────
# B. 考勤登记
# ─────────────────────────────────────────────
elif menu == "考勤登记":
    st.title("✅ 考勤与进度记录")
    if not students_db:
        st.info("暂无学生资料。")
    else:
        sid = st.selectbox("选择学生", list(students_db.keys()),
                           format_func=lambda x: f"{students_db[x]['name']} ({x})")
        with st.form("attendance_form", clear_on_submit=True):
            status_option = st.radio(
                "课堂状态",
                ["Attended","Late Notice","Cancelled (Student)","Cancelled (Teacher)","Temporary Stop"],
                horizontal=True,
            )
            date_val = st.date_input("日期", datetime.now())
            remarks  = st.text_area("今日学习进度/评语")

            # Show makeup minutes field only for cancelled
            cancel_minutes = st.number_input(
                "需补课时长（分钟）", min_value=0, max_value=180, value=60, step=15,
                help="仅在 Cancelled 时有效"
            )

            if st.form_submit_button("💾 保存考勤记录"):
                if "Cancelled" in status_option:
                    final_remarks = f"N/A - {status_option}"
                    students_db[sid]["replacement_credits"] += 1
                    students_db[sid]["replacement_minutes"] = students_db[sid].get("replacement_minutes", 0) + cancel_minutes
                elif status_option == "Temporary Stop":
                    final_remarks = remarks or "Temporary Stop – 无需补课"
                else:
                    final_remarks = remarks
                students_db[sid]["history"].append({
                    "date":           str(date_val),
                    "status":         status_option,
                    "remarks":        final_remarks,
                    "replaced":       False,
                    "cancel_minutes": cancel_minutes if "Cancelled" in status_option else 0,
                })
                save_data(students_db)
                st.success("✅ 记录已同步到云端！")

# ─────────────────────────────────────────────
# C. 补课安排
# ─────────────────────────────────────────────
elif menu == "补课安排":
    st.title("🔄 补课排课")
    if not students_db:
        st.info("暂无学生资料。")
    else:
        sid = st.selectbox("选择学生", list(students_db.keys()),
                           format_func=lambda x: f"{students_db[x]['name']} ({x})")
        history = students_db[sid]["history"]
        cancelled_indices = [
            i for i, log in enumerate(history)
            if "Cancelled" in log.get("status","") and not log.get("replaced")
        ]
        if cancelled_indices:
            selection = st.selectbox(
                "选择欠课", cancelled_indices,
                format_func=lambda x: f"{history[x]['date']} ({history[x]['status']}) – 欠 {history[x].get('cancel_minutes',0)} 分钟",
            )
            with st.form("rep_form"):
                rep_type    = st.radio("补课类型", ["Fully Replace","Part Replace"], horizontal=True)
                rep_date    = st.date_input("补课日期")
                rep_time    = st.time_input("补课上课时间")
                rep_minutes = st.number_input(
                    "本次补课时长（分钟）", min_value=15, max_value=180, value=30, step=15,
                    help="Fully Replace 时自动使用原欠课时长"
                )
                rep_remarks = st.text_area("补课进度")
                if st.form_submit_button("确认补课"):
                    orig_minutes = history[selection].get('cancel_minutes', 0)
                    if rep_type == "Fully Replace":
                        students_db[sid]["history"][selection]["replaced"] = True
                        students_db[sid]["replacement_credits"] -= 1
                        students_db[sid]["history"][selection]["cancel_minutes"] = 0
                        status_label = "Fully Replaced"
                        done_min = orig_minutes
                    else:
                        remaining = orig_minutes - rep_minutes
                        if remaining <= 0:
                            students_db[sid]["history"][selection]["replaced"] = True
                            students_db[sid]["replacement_credits"] -= 1
                            students_db[sid]["history"][selection]["cancel_minutes"] = 0
                        else:
                            students_db[sid]["history"][selection]["cancel_minutes"] = int(remaining)
                        status_label = "Part Replaced"
                        done_min = rep_minutes
                    students_db[sid]["history"].append({
                        "date": str(rep_date),
                        "status": status_label,
                        "remarks": f"[{rep_time.strftime('%H:%M')}] {rep_remarks}",
                        "replaced": True,
                        "cancel_minutes": done_min,
                    })
                    save_data(students_db)
                    st.success(f"✅ {rep_type} 已记录！（{done_min} 分钟）")
                    st.rerun()
        else:
            st.info("没有待补课记录。")

# ─────────────────────────────────────────────
# D. 报告查询
# ─────────────────────────────────────────────
elif menu == "报告查询":
    st.title("📊 学生档案卡")
    if not students_db:
        st.info("暂无学生资料。")
    else:
        sid  = st.selectbox("选择学生", list(students_db.keys()),
                            format_func=lambda x: f"{students_db[x]['name']} ({x})")
        info = students_db[sid]
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("乐器", info['instrument'])
        m2.metric("等级", info['grade'])
        # Recalculate from history to ensure accuracy
        pending_min = sum(
            h.get('cancel_minutes', 0) for h in info['history']
            if "Cancelled" in h.get('status','') and not h.get('replaced')
        )
        pending_count = sum(
            1 for h in info['history']
            if "Cancelled" in h.get('status','') and not h.get('replaced')
        )
        m3.metric("待补课次数", pending_count)
        m4.metric("待补课时长", f"{pending_min} min")

        if info.get('phone') or info.get('email'):
            st.write(f"📞 {info.get('phone','')}　✉️ {info.get('email','')}")

        if info.get('schedule'):
            for s in info['schedule']:
                st.write(f"📅 {s['day']}　{s.get('start','')[:5]} – {s.get('end','')[:5]}　（{s.get('duration_min','')} 分钟）")

        st.divider()
        st.subheader("📜 历史进度表")
        if info['history']:
            if role == "viewer":
                st.dataframe(info['history'], use_container_width=True)
            else:
                edited = st.data_editor(info['history'], num_rows="dynamic", use_container_width=True)
                if st.button("💾 保存表格修改"):
                    students_db[sid]['history'] = edited
                    save_data(students_db)
                    st.success("✅ 已更新！")
                    time.sleep(1)
                    st.rerun()
        else:
            st.write("尚无考勤历史。")

        st.divider()
        st.subheader("🤖 AI 学习进度分析")
        lang = st.radio("报告语言", ["English","中文"], horizontal=True)
        if st.button("✨ 生成 AI 分析"):
            with st.spinner("AI 分析中..."):
                result = call_gemini(build_prompt(sid, info, lang))
            st.session_state[f'ai_result_{sid}'] = result

        if f'ai_result_{sid}' in st.session_state:
            result = st.session_state[f'ai_result_{sid}']
            st.info(result)

            # WhatsApp button
            phone = info.get('phone', '').strip().replace(' ', '').replace('-', '')
            if phone:
                # Format Malaysian number
                if phone.startswith('0'):
                    phone = '6' + phone
                msg = f"Hi, here is today's lesson report for {info['name']}:%0A%0A{result.replace(chr(10), '%0A')}"
                wa_url = f"https://wa.me/{phone}?text={msg}"
                st.markdown(f"[📲 发送给家长 (WhatsApp)]({wa_url})", unsafe_allow_html=True)
            else:
                st.caption("⚠️ 该学生未填写联络号码，无法发送 WhatsApp")

# ─────────────────────────────────────────────
# E. 管理学生信息
# ─────────────────────────────────────────────
elif menu == "管理学生信息":
    st.title("⚙️ 学生信息维护")
    if not students_db:
        st.info("暂无学生资料。")
    else:
        sid = st.selectbox("修改学生", list(students_db.keys()),
                           format_func=lambda x: f"{students_db[x]['name']} ({x})")
        col_edit, col_del = st.columns(2)
        with col_edit:
            st.subheader("修改资料")
            new_id    = st.text_input("修改学生 ID", value=sid)
            new_name  = st.text_input("修改姓名",    value=students_db[sid]['name'])
            new_inst  = st.selectbox("修改乐器", INSTRUMENTS,
                                     index=INSTRUMENTS.index(students_db[sid]['instrument'])
                                     if students_db[sid]['instrument'] in INSTRUMENTS else 0)
            new_grade = st.text_input("修改等级", value=students_db[sid]['grade'])
            new_phone = st.text_input("修改联络号码", value=students_db[sid].get('phone',''))
            new_email = st.text_input("修改电子邮箱", value=students_db[sid].get('email',''))

            st.write("**修改固定上课时间**")
            current_schedule = students_db[sid].get('schedule', [])
            num_slots = st.number_input("每周上课次数", min_value=1, max_value=7,
                                        value=max(1, len(current_schedule)), step=1)
            new_duration = st.number_input("每节课时长（分钟）", min_value=15, max_value=180,
                                           value=current_schedule[0].get('duration_min', 60) if current_schedule else 60,
                                           step=15)
            new_schedule = []
            for i in range(int(num_slots)):
                c1, c2 = st.columns(2)
                ex_day   = current_schedule[i].get('day', DAYS_OPTIONS[0])  if i < len(current_schedule) else DAYS_OPTIONS[0]
                ex_start = current_schedule[i].get('start', "09:00:00") if i < len(current_schedule) else "09:00:00"
                with c1:
                    day = st.selectbox(f"第{i+1}节 – 周几", DAYS_OPTIONS,
                                       index=DAYS_OPTIONS.index(ex_day) if ex_day in DAYS_OPTIONS else 0,
                                       key=f"ed_{i}")
                with c2:
                    tp = ex_start.split(":")
                    start_t = st.time_input(f"第{i+1}节 – 开始时间",
                                            value=dtime(int(tp[0]), int(tp[1])), key=f"et_{i}")
                start_dt = datetime.combine(date.today(), start_t)
                end_dt   = start_dt + timedelta(minutes=int(new_duration))
                end_t    = end_dt.time()
                st.caption(f"↳ 结束时间：{end_t.strftime('%H:%M')}")
                new_schedule.append({
                    "day":          day,
                    "start":        str(start_t),
                    "end":          end_t.strftime("%H:%M:%S"),
                    "duration_min": int(new_duration),
                })

            if st.button("更新学生资料"):
                info = students_db.pop(sid)
                info.update({
                    "name": new_name, "instrument": new_inst,
                    "grade": new_grade, "phone": new_phone,
                    "email": new_email, "schedule": new_schedule,
                })
                students_db[new_id] = info
                save_data(students_db)
                st.success("✅ 更新成功！")
                st.rerun()

        with col_del:
            st.subheader("删除学生")
            st.warning("⚠️ 此操作不可撤销")
            if st.button("🚨 彻底删除学生"):
                del students_db[sid]
                save_data(students_db)
                st.rerun()

# ─────────────────────────────────────────────
# F. 申请修改学生信息 (editor)
# ─────────────────────────────────────────────
elif menu == "申请修改学生信息":
    st.title("📝 申请修改学生信息")
    st.info("ℹ️ 修改申请需要 Owner 批准后才会生效。")
    if students_db:
        sid       = st.selectbox("选择学生", list(students_db.keys()),
                                 format_func=lambda x: f"{students_db[x]['name']} ({x})")
        new_name  = st.text_input("修改姓名",    value=students_db[sid]['name'])
        new_grade = st.text_input("修改等级",    value=students_db[sid]['grade'])
        new_inst  = st.selectbox("修改乐器", INSTRUMENTS,
                                 index=INSTRUMENTS.index(students_db[sid]['instrument'])
                                 if students_db[sid]['instrument'] in INSTRUMENTS else 0)
        new_phone = st.text_input("修改联络号码", value=students_db[sid].get('phone',''))
        new_email = st.text_input("修改电子邮箱", value=students_db[sid].get('email',''))
        if st.button("📤 提交修改申请"):
            submit_pending("edit_student", {
                "student_id": sid,
                "new_name":  new_name,  "old_name":  students_db[sid]['name'],
                "new_grade": new_grade, "old_grade": students_db[sid]['grade'],
                "new_inst":  new_inst,  "old_inst":  students_db[sid]['instrument'],
                "new_phone": new_phone, "new_email": new_email,
            })
            st.success("✅ 申请已提交，等待 Owner 批准！")

# ─────────────────────────────────────────────
# G. 报表中心
# ─────────────────────────────────────────────
elif menu == "报表中心":
    st.title("📋 报表中心")

    st.subheader("⏳ 全部待补课汇总")
    rows = []
    for s, i in students_db.items():
        if i['replacement_credits'] > 0 or i.get('replacement_minutes', 0) > 0:
            pending_min = sum(
                h.get('cancel_minutes', 0) for h in i['history']
                if "Cancelled" in h.get('status','') and not h.get('replaced')
            )
            pending_count = sum(
                1 for h in i['history']
                if "Cancelled" in h.get('status','') and not h.get('replaced')
            )
            rows.append({
                "学生ID":    s,
                "姓名":      i['name'],
                "乐器":      i['instrument'],
                "待补课次数": pending_count,
                "待补课时长(min)": pending_min,
            })
    if rows:
        df_p = pd.DataFrame(rows).sort_values("待补课次数", ascending=False)
        col1, col2 = st.columns(2)
        col1.metric("全部待补课总次数", df_p["待补课次数"].sum())
        col2.metric("全部待补课总时长", f"{df_p['待补课时长(min)'].sum()} min")
        st.dataframe(df_p, use_container_width=True)
    else:
        st.success("✅ 目前没有待补课记录！")

    st.divider()

    st.subheader("🤖 AI 学生报告导出")
    report_mode = st.radio("生成对象", ["单个学生","全部学生"], horizontal=True)
    lang        = st.radio("报告语言", ["English","中文"], horizontal=True, key="rl")

    if report_mode == "单个学生":
        sid = st.selectbox("选择学生", list(students_db.keys()),
                           format_func=lambda x: f"{students_db[x]['name']} ({x})")
        if st.button("✨ 生成报告"):
            with st.spinner("AI 生成中..."):
                report = call_gemini(build_prompt(sid, students_db[sid], lang))
            st.session_state['report']      = report
            st.session_state['report_name'] = students_db[sid]['name']
    else:
        if st.button("✨ 生成全部学生报告"):
            with st.spinner("AI 生成中（可能需要较长时间）..."):
                parts    = []
                progress = st.progress(0)
                total_s  = len(students_db)
                for idx, (sid, info) in enumerate(students_db.items()):
                    r = call_gemini(build_prompt(sid, info, lang))
                    parts.append(f"{'='*50}\n{info['name']} ({sid})\n{'='*50}\n{r}\n")
                    progress.progress((idx + 1) / total_s)
                report = "\n".join(parts)
            st.session_state['report']      = report
            st.session_state['report_name'] = "All_Students"

    if 'report' in st.session_state:
        st.text_area("📄 报告内容", st.session_state['report'], height=300)
        fname = f"{st.session_state['report_name']}_report_{datetime.now().strftime('%Y%m%d')}"
        col1, col2 = st.columns(2)
        with col1:
            st.download_button("⬇️ 下载 TXT", data=st.session_state['report'].encode("utf-8"),
                               file_name=f"{fname}.txt", mime="text/plain")
        with col2:
            try:
                from docx import Document
                from io import BytesIO
                doc = Document()
                doc.add_heading("Brent Music – Student Report", 0)
                doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
                doc.add_paragraph("")
                for line in st.session_state['report'].split("\n"):
                    doc.add_paragraph(line)
                buf = BytesIO(); doc.save(buf); buf.seek(0)
                st.download_button("⬇️ 下载 Word 文档", data=buf,
                                   file_name=f"{fname}.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except ImportError:
                st.info("💡 Word 导出需在 requirements.txt 加入 `python-docx`")

# ─────────────────────────────────────────────
# H. 课程日历
# ─────────────────────────────────────────────
elif menu == "课程日历":
    st.title("📅 课程日历")

    weekly = {i: [] for i in range(7)}
    for sid, info in students_db.items():
        for slot in info.get('schedule', []):
            d = DAY_MAP.get(slot['day'], -1)
            if d >= 0:
                weekly[d].append({
                    "start": slot.get('start',''), "end": slot.get('end',''),
                    "name": info['name'], "instrument": info['instrument'], "sid": sid,
                })

    st.subheader("📆 本周固定课程表")
    cols = st.columns(7)
    for idx, col in enumerate(cols):
        with col:
            st.markdown(f"**{DAYS_OPTIONS[idx]}**")
            slots = sorted(weekly[idx], key=lambda x: x['start'])
            if slots:
                for s in slots:
                    st.markdown(
                        f"<div style='background:#1a3a5c;border-radius:8px;padding:8px;"
                        f"margin-bottom:6px;font-size:12px;color:white;border-left:3px solid #4a9eff'>"
                        f"⏰ {s['start'][:5]}–{s['end'][:5]}<br>"
                        f"<b>{s['name']}</b><br><i>{s['instrument']}</i></div>",
                        unsafe_allow_html=True,
                    )
            else:
                st.caption("无课")

    st.divider()
    st.subheader("➕ 添加单次临时课程")
    with st.form("oneoff"):
        c1, c2, c3 = st.columns(3)
        with c1:
            oo_sid  = st.selectbox("学生", list(students_db.keys()),
                                   format_func=lambda x: f"{students_db[x]['name']} ({x})")
        with c2:
            oo_date = st.date_input("日期", datetime.now())
        with c3:
            oo_time = st.time_input("时间")
        oo_note = st.text_input("备注（可选）")
        if st.form_submit_button("✅ 添加临时课程"):
            students_db[oo_sid]["history"].append({
                "date": str(oo_date), "status": "Scheduled (One-off)",
                "remarks": oo_note or "临时排课", "replaced": False, "cancel_minutes": 0,
            })
            save_data(students_db)
            st.success("✅ 已添加！")
            st.rerun()

    st.divider()
    st.subheader("🔗 导出到 Google Calendar")
    today = date.today()

    def next_weekday(idx):
        ahead = idx - today.weekday()
        if ahead <= 0:
            ahead += 7
        return today + timedelta(days=ahead)

    for sid, info in students_db.items():
        if info.get('schedule'):
            with st.expander(f"📅 {info['name']} ({info['instrument']})"):
                for slot in info['schedule']:
                    d_idx = DAY_MAP.get(slot['day'], 0)
                    ld    = next_weekday(d_idx)
                    tp    = slot.get('start','09:00:00').split(":")
                    h, m  = int(tp[0]), int(tp[1])
                    tp2   = slot.get('end','10:00:00').split(":")
                    h2, m2 = int(tp2[0]), int(tp2[1])
                    start = ld.strftime("%Y%m%d") + f"T{h:02d}{m:02d}00"
                    end   = ld.strftime("%Y%m%d") + f"T{h2:02d}{m2:02d}00"
                    title = f"{info['name']}+–+{info['instrument']}+Lesson"
                    url   = (f"https://calendar.google.com/calendar/render?action=TEMPLATE"
                             f"&text={title}&dates={start}/{end}"
                             f"&recur=RRULE:FREQ%3DWEEKLY&details=Brent+Music+Lesson")
                    st.markdown(f"[📆 {slot['day']} {slot.get('start','')[:5]}–{slot.get('end','')[:5]} → 添加到 Google Calendar]({url})")

# ─────────────────────────────────────────────
# I. To Do List
# ─────────────────────────────────────────────
elif menu == "📝 To Do List":
    st.title("📝 To Do List")

    todos = load_todos()

    # Add new task
    with st.form("add_todo", clear_on_submit=True):
        new_task = st.text_input("新增任务")
        if st.form_submit_button("➕ 添加"):
            if new_task.strip():
                todos.append({
                    "id":         str(int(time.time())),
                    "task":       new_task.strip(),
                    "done":       False,
                    "created_by": username,
                    "timestamp":  datetime.now().strftime("%Y-%m-%d %H:%M"),
                })
                save_todos(todos)
                st.rerun()

    st.divider()

    if not todos:
        st.info("暂无任务，添加一个吧！")
    else:
        pending_todos   = [t for t in todos if not t['done']]
        completed_todos = [t for t in todos if t['done']]

        if pending_todos:
            st.subheader("⏳ 待完成")
            for t in pending_todos:
                col1, col2 = st.columns([8, 1])
                with col1:
                    st.write(f"☐ {t['task']}  \n<small style='color:gray'>by {t['created_by']} · {t['timestamp']}</small>", unsafe_allow_html=True)
                with col2:
                    if st.button("✅", key=f"done_{t['id']}"):
                        todos = [x for x in todos if x['id'] != t['id']]
                        save_todos(todos)
                        st.rerun()

        if completed_todos:
            st.divider()
            st.subheader("✅ 已完成")
            for t in completed_todos:
                col1, col2 = st.columns([8, 1])
                with col1:
                    st.markdown(f"<s style='color:gray'>{t['task']}</s>　<small style='color:gray'>by {t['created_by']}</small>", unsafe_allow_html=True)
                with col2:
                    if st.button("🗑️", key=f"del_{t['id']}"):
                        todos = [x for x in todos if x['id'] != t['id']]
                        save_todos(todos)
                        st.rerun()

# ─────────────────────────────────────────────
# J. Events
# ─────────────────────────────────────────────
elif menu == "🎪 Events":
    st.title("🎪 Events（考试 / Recital）")

    events = load_events()

    # Add new event
    with st.form("add_event", clear_on_submit=True):
        st.subheader("➕ 新增 Event")
        ec1, ec2 = st.columns(2)
        with ec1:
            ev_sid  = st.selectbox("学生", list(students_db.keys()),
                                   format_func=lambda x: f"{students_db[x]['name']} ({x})")
            ev_type = st.selectbox("类型", ["考试 (Exam)","演奏会 (Recital)","比赛 (Competition)","其他"])
        with ec2:
            ev_name = st.text_input("Event 名称")
            ev_date = st.date_input("日期", datetime.now())
        ev_notes = st.text_area("备注")
        if st.form_submit_button("✅ 添加 Event"):
            if ev_name:
                events.append({
                    "id":           str(int(time.time())),
                    "student_id":   ev_sid,
                    "student_name": students_db[ev_sid]['name'],
                    "event_type":   ev_type,
                    "event_name":   ev_name,
                    "event_date":   str(ev_date),
                    "notes":        ev_notes,
                    "created_by":   username,
                })
                save_events(events)
                st.success("✅ Event 已添加！")
                st.rerun()
            else:
                st.error("请填入 Event 名称")

    st.divider()

    if not events:
        st.info("暂无 Event 记录。")
    else:
        # Sort by date
        events_sorted = sorted(events, key=lambda x: x['event_date'])

        # Upcoming
        today_str = str(date.today())
        upcoming  = [e for e in events_sorted if e['event_date'] >= today_str]
        past      = [e for e in events_sorted if e['event_date'] < today_str]

        if upcoming:
            st.subheader("📅 即将到来")
            for e in upcoming:
                days_left = (date.fromisoformat(e['event_date']) - date.today()).days
                with st.expander(f"🔔 {e['event_date']} – {e['student_name']} – {e['event_name']} （{days_left} 天后）"):
                    st.write(f"**类型：** {e['event_type']}")
                    st.write(f"**学生：** {e['student_name']} ({e['student_id']})")
                    st.write(f"**备注：** {e['notes'] or '–'}")
                    if st.button("🗑️ 删除", key=f"del_ev_{e['id']}"):
                        events = [x for x in events if x['id'] != e['id']]
                        save_events(events)
                        st.rerun()

        if past:
            st.divider()
            st.subheader("📁 过去的 Events")
            for e in reversed(past):
                with st.expander(f"📌 {e['event_date']} – {e['student_name']} – {e['event_name']}"):
                    st.write(f"**类型：** {e['event_type']}")
                    st.write(f"**备注：** {e['notes'] or '–'}")
                    if st.button("🗑️ 删除", key=f"del_pev_{e['id']}"):
                        events = [x for x in events if x['id'] != e['id']]
                        save_events(events)
                        st.rerun()

# ─────────────────────────────────────────────
# K. 审批请求 (owner)
# ─────────────────────────────────────────────
elif menu == "⚠️ 审批请求":
    st.title("⚠️ 待审批的修改请求")
    pending = load_pending()
    if not pending:
        st.success("✅ 目前没有待审批的请求。")
    else:
        for p in pending:
            with st.expander(f"📋 {p['timestamp']} — {p['requested_by']} 申请修改"):
                data = p['data']
                st.write(f"**学生 ID：** {data['student_id']}")
                st.write(f"**姓名：** {data.get('old_name','–')} → {data.get('new_name','–')}")
                st.write(f"**等级：** {data.get('old_grade','–')} → {data.get('new_grade','–')}")
                st.write(f"**乐器：** {data.get('old_inst','–')} → {data.get('new_inst','–')}")
                c1, c2 = st.columns(2)
                with c1:
                    if st.button("✅ 批准", key=f"a_{p['id']}"):
                        sid = data['student_id']
                        if sid in students_db:
                            students_db[sid].update({
                                'name':       data.get('new_name',  students_db[sid]['name']),
                                'grade':      data.get('new_grade', students_db[sid]['grade']),
                                'instrument': data.get('new_inst',  students_db[sid]['instrument']),
                                'phone':      data.get('new_phone', students_db[sid].get('phone','')),
                                'email':      data.get('new_email', students_db[sid].get('email','')),
                            })
                            save_data(students_db)
                        pending = [x for x in pending if x['id'] != p['id']]
                        save_pending(pending)
                        st.success("✅ 已批准！")
                        st.rerun()
                with c2:
                    if st.button("❌ 拒绝", key=f"r_{p['id']}"):
                        pending = [x for x in pending if x['id'] != p['id']]
                        save_pending(pending)
                        st.warning("已拒绝。")
                        st.rerun()
